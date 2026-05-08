from docx import Document
from docx.shared import Pt, RGBColor, Inches, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_ALIGN_VERTICAL
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import copy

doc = Document()

# ── 페이지 여백 설정 ─────────────────────────────────────────────────
for section in doc.sections:
    section.top_margin    = Cm(2.0)
    section.bottom_margin = Cm(2.0)
    section.left_margin   = Cm(2.5)
    section.right_margin  = Cm(2.5)

# ── 색상 ─────────────────────────────────────────────────────────────
NAVY   = RGBColor(0x0b, 0x1e, 0x45)
BLUE   = RGBColor(0x3B, 0x5A, 0xC6)
RED    = RGBColor(0xd4, 0x00, 0x2a)
WHITE  = RGBColor(0xFF, 0xFF, 0xFF)
LIGHT  = RGBColor(0xF4, 0xF5, 0xF7)
GRAY   = RGBColor(0x55, 0x55, 0x55)
GREEN  = RGBColor(0x1D, 0x7A, 0x4A)
YELLOW = RGBColor(0xb4, 0x53, 0x09)
PURPLE = RGBColor(0x6d, 0x28, 0xd9)

# ── 헬퍼 ─────────────────────────────────────────────────────────────
def set_cell_bg(cell, color: RGBColor):
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    shd = OxmlElement('w:shd')
    hex_color = f"{color[0]:02X}{color[1]:02X}{color[2]:02X}"
    shd.set(qn('w:val'), 'clear')
    shd.set(qn('w:color'), 'auto')
    shd.set(qn('w:fill'), hex_color)
    tcPr.append(shd)

def set_para_bg(para, color: RGBColor):
    pPr = para._p.get_or_add_pPr()
    shd = OxmlElement('w:shd')
    hex_color = f"{color[0]:02X}{color[1]:02X}{color[2]:02X}"
    shd.set(qn('w:val'), 'clear')
    shd.set(qn('w:color'), 'auto')
    shd.set(qn('w:fill'), hex_color)
    pPr.append(shd)

def heading(text, level=1, color=NAVY, size=22):
    p = doc.add_paragraph()
    set_para_bg(p, NAVY if level==1 else BLUE if level==2 else LIGHT)
    run = p.add_run("  " + text)
    run.font.size = Pt(size)
    run.font.bold = True
    run.font.color.rgb = WHITE if level <= 2 else NAVY
    p.paragraph_format.space_before = Pt(16 if level==1 else 10)
    p.paragraph_format.space_after  = Pt(6)
    p.paragraph_format.left_indent  = Inches(0)
    return p

def body(text, indent=0, color=None, size=11, bold=False, space_before=2):
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.font.size = Pt(size)
    run.font.bold = bold
    if color:
        run.font.color.rgb = color
    p.paragraph_format.space_before = Pt(space_before)
    p.paragraph_format.space_after  = Pt(2)
    if indent:
        p.paragraph_format.left_indent = Inches(indent)
    return p

def bullet(text, level=0, color=None):
    p = doc.add_paragraph(style='List Bullet')
    run = p.add_run(text)
    run.font.size = Pt(10.5)
    if color:
        run.font.color.rgb = color
    p.paragraph_format.left_indent = Inches(0.3 + level * 0.2)
    p.paragraph_format.space_before = Pt(2)
    p.paragraph_format.space_after  = Pt(2)
    return p

def add_table(headers, rows, col_widths=None, header_color=NAVY):
    t = doc.add_table(rows=1+len(rows), cols=len(headers))
    t.style = 'Table Grid'
    # 헤더
    hdr = t.rows[0]
    for i, h in enumerate(headers):
        cell = hdr.cells[i]
        set_cell_bg(cell, header_color)
        p = cell.paragraphs[0]
        run = p.add_run(h)
        run.font.bold = True
        run.font.size = Pt(10)
        run.font.color.rgb = WHITE
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    # 데이터
    for ri, row in enumerate(rows):
        for ci, val in enumerate(row):
            cell = t.rows[ri+1].cells[ci]
            if ri % 2 == 0:
                set_cell_bg(cell, RGBColor(0xF0,0xF4,0xFF))
            p = cell.paragraphs[0]
            run = p.add_run(str(val))
            run.font.size = Pt(9.5)
            if ci == 0:
                run.font.bold = True
                run.font.color.rgb = BLUE
    # 열 너비
    if col_widths:
        for row in t.rows:
            for i, w in enumerate(col_widths):
                row.cells[i].width = Inches(w)
    doc.add_paragraph()
    return t

# ════════════════════════════════════════════════════════════════════
# 표지
# ════════════════════════════════════════════════════════════════════
p = doc.add_paragraph()
set_para_bg(p, NAVY)
run = p.add_run("  🐢  SlowBro — 기술 문서")
run.font.size = Pt(32)
run.font.bold = True
run.font.color.rgb = WHITE
p.paragraph_format.space_before = Pt(20)
p.paragraph_format.space_after  = Pt(4)

p2 = doc.add_paragraph()
set_para_bg(p2, NAVY)
run2 = p2.add_run("  어르신을 위한 접근성 티켓 예매 브라우저")
run2.font.size = Pt(16)
run2.font.color.rgb = RGBColor(0xBB, 0xCC, 0xFF)
p2.paragraph_format.space_after = Pt(4)

p3 = doc.add_paragraph()
set_para_bg(p3, NAVY)
run3 = p3.add_run('  "천천히 해도 괜찮아요 🌿"')
run3.font.size = Pt(13)
run3.font.italic = True
run3.font.color.rgb = RGBColor(0x88, 0xAA, 0xFF)
p3.paragraph_format.space_after = Pt(20)

body("Electron 28  ·  Cheerio  ·  Web Speech API  ·  BrowserView  ·  Vanilla JS",
     color=GRAY, size=10)

doc.add_paragraph()

# ════════════════════════════════════════════════════════════════════
# 1. 문제 정의
# ════════════════════════════════════════════════════════════════════
heading("01  문제 정의", level=1)
body("기존 티켓 예매 사이트는 어르신에게 다음과 같은 장벽을 만듭니다:", size=11)

problems = [
    ("⚡ 속도 압박",   "30초 내 좌석 선택 강제, 팝업·광고 난무"),
    ("👀 정보 과부하", "작은 글씨, 수십 개 버튼, 어디를 눌러야 할지 모름"),
    ("❌ 예매 실패",   "에러 메시지 이해 불가, 시간 초과 반복"),
]
t = doc.add_table(rows=1+len(problems), cols=2)
t.style = 'Table Grid'
hdr = t.rows[0]
for i, h in enumerate(["문제", "설명"]):
    cell = hdr.cells[i]
    set_cell_bg(cell, RED)
    run = cell.paragraphs[0].add_run(h)
    run.font.bold = True; run.font.size = Pt(10); run.font.color.rgb = WHITE
    cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
for ri, (prob, desc) in enumerate(problems):
    t.rows[ri+1].cells[0].paragraphs[0].add_run(prob).font.bold = True
    t.rows[ri+1].cells[0].paragraphs[0].runs[0].font.size = Pt(10)
    set_cell_bg(t.rows[ri+1].cells[0], RGBColor(0xF0,0xF4,0xFF))
    t.rows[ri+1].cells[1].paragraphs[0].add_run(desc).font.size = Pt(10)
    t.rows[ri+1].cells[0].width = Inches(2.0)
    t.rows[ri+1].cells[1].width = Inches(10.0)
doc.add_paragraph()

heading("SlowBro 솔루션", level=2, size=14)
solutions = [
    "🐢  핵심 단계만 추려 1단계씩 안내 (Step-by-step guide)",
    "🎨  원본 사이트 디자인을 계승하며 불필요 요소 제거 (SITE_THEMES)",
    "🔊  한국어 TTS로 각 단계 읽어주기 (Web Speech API ko-KR)",
    "🛡️  이중 클릭·도배 방지 Guard로 실수 예매 차단 (antiAbuse.js)",
    "🌐  원본 사이트도 그대로 볼 수 있는 BrowserView 모드 제공",
]
for s in solutions:
    bullet(s)

doc.add_paragraph()

# ════════════════════════════════════════════════════════════════════
# 2. Tech Stack
# ════════════════════════════════════════════════════════════════════
heading("02  Tech Stack", level=1)

stacks = [
    ("Electron 28",       "데스크톱 앱",  "Main/Renderer 이중 프로세스, contextBridge IPC 보안"),
    ("Cheerio",           "파싱 엔진",   "서버사이드 jQuery — 실제 사이트 HTML에서 날짜/좌석/버튼 추출"),
    ("Web Speech API",    "TTS",        "브라우저 내장 한국어 음성 합성 (ko-KR, 0.9× 속도)"),
    ("BrowserView",       "임베드",      "Electron 내장 Chromium — 원본 사이트 그대로 앱 안에 표시"),
    ("CSS Custom Props",  "다이나믹 테마","--site-primary/accent/bg 로 원본 사이트 색상 동적 적용"),
    ("electron-builder",  "패키징",      "Windows(NSIS) / Mac / Linux 인스톨러 자동 생성"),
    ("Vanilla JS",        "렌더러",      "외부 프레임워크 없이 순수 JS — 상태 관리·이벤트·TTS 처리"),
]
add_table(
    ["기술", "역할", "설명"],
    stacks,
    col_widths=[1.8, 1.2, 9.0],
    header_color=BLUE
)

heading("레이어별 파일 구조", level=2, size=14)
layers = [
    ("src/main/index.js",     "Main Process",    "BrowserWindow, BrowserView, IPC 핸들러"),
    ("src/preload/index.js",  "contextBridge",   "Main↔Renderer 안전한 API 노출"),
    ("src/engine/pageParser.js","파싱 엔진",     "SITE_RULES + Cheerio + SITE_THEMES"),
    ("src/guard/antiAbuse.js","예매 가드",       "클릭 속도·세션 중복 체크"),
    ("renderer/index.html",   "렌더러 UI",       "앱 레이아웃 및 모드 토글"),
    ("renderer/pages/app.js", "렌더러 로직",     "상태 관리, TTS, 테마, BrowserView 제어"),
    ("renderer/styles/main.css","스타일",        "CSS Custom Properties 기반 동적 테마"),
    ("mock-sites/",           "목 사이트",       "Giants·Interpark·Korail 로컬 HTML 모의 사이트"),
]
add_table(
    ["파일 경로", "레이어", "역할"],
    layers,
    col_widths=[3.5, 2.0, 6.5],
    header_color=NAVY
)

# ════════════════════════════════════════════════════════════════════
# 3. 아키텍처
# ════════════════════════════════════════════════════════════════════
heading("03  아키텍처", level=1)

body("Electron의 Main/Renderer 이중 프로세스 구조를 기반으로 contextBridge IPC 통신으로 분리합니다.", size=11)
doc.add_paragraph()

arch_text = """[Main Process]                    IPC (contextBridge)          [Renderer Process]
  index.js (Window·BrowserView)  ←── invoke/send/on ───►   index.html + app.js
  pageParser.js (Cheerio)                                   Web Speech API (TTS)
  antiAbuse.js (Guard)                                      CSS Custom Properties
  preload/index.js                                          (동적 사이트 테마)
           │
           └─► [BrowserView]  ── 원본 사이트 직접 임베드 ──►  Chromium 렌더링"""

p = doc.add_paragraph()
run = p.add_run(arch_text)
run.font.name = 'Courier New'
run.font.size = Pt(9)
set_para_bg(p, RGBColor(0x1a, 0x2a, 0x55))
run.font.color.rgb = RGBColor(0xBB, 0xDD, 0xFF)
p.paragraph_format.space_before = Pt(4)
p.paragraph_format.space_after  = Pt(8)

heading("IPC 통신 원칙", level=2, size=14)
ipc_principles = [
    "모든 Main-Renderer 통신은 ipcRenderer.invoke (양방향) 또는 ipcRenderer.send (단방향) 사용",
    "contextIsolation: true — nodeIntegration: false 로 보안 강화",
    "BrowserView 이벤트(URL 변경, 로딩 상태)는 mainWindow.webContents.send() 로 푸시",
    "Renderer는 window.api (contextBridge 노출 객체)를 통해서만 Main 기능 접근 가능",
]
for p in ipc_principles:
    bullet(p)

doc.add_paragraph()

# ════════════════════════════════════════════════════════════════════
# 4. 시스템 플로우
# ════════════════════════════════════════════════════════════════════
heading("04  시스템 플로우", level=1)

heading("Flow 1 — SlowBro 변환 모드", level=2, size=14)
flow1 = [
    ("1", "URL 입력",          "url-input에 mock:// 또는 https:// URL 입력 후 이동 버튼"),
    ("2", "page:load 호출",    "Renderer → Main: api.loadPage(url) 으로 IPC invoke"),
    ("3", "Cheerio 파싱",      "SITE_RULES 셀렉터로 날짜·좌석·버튼 추출, SITE_THEMES 반환"),
    ("4", "Steps + Theme 반환","{ steps[], theme } 응답 → Renderer 수신"),
    ("5", "applyTheme()",      "CSS --site-primary/accent/bg 변수 설정, 로고 교체"),
    ("6", "renderStep() 반복", "단계별 카드 표시 → TTS 읽기 → 다음 단계 진행"),
]
add_table(
    ["단계", "동작", "세부 설명"],
    flow1,
    col_widths=[0.6, 2.5, 8.9],
    header_color=BLUE
)

heading("Flow 2 — 원본 브라우저 모드", level=2, size=14)
flow2 = [
    ("1", "원본 보기 클릭",   "mode-toggle 버튼 → switchMode('orig') 호출"),
    ("2", "browser:open",    "api.browser.open(url) IPC invoke → BrowserView 생성"),
    ("3", "setBounds 계산",  "navbar.getBoundingClientRect().bottom 기준 높이 계산"),
    ("4", "URL 로드",        "browserView.webContents.loadURL(url)"),
    ("5", "이벤트 수신",     "did-navigate, loading → mainWindow.webContents.send() 푸시"),
    ("6", "네비게이션 바 업데이트", "bv-url-display, bv-loading 실시간 반영"),
]
add_table(
    ["단계", "동작", "세부 설명"],
    flow2,
    col_widths=[0.6, 2.5, 8.9],
    header_color=RGBColor(0x0b,0x6e,0x9e)
)

heading("Flow 3 — 예매 가드 (antiAbuse)", level=2, size=14)
flow3 = [
    ("1", "예매 버튼 클릭",   "renderStep에서 선택 확정 버튼 이벤트 발생"),
    ("2", "checkClick",      "guard:checkClick 호출 — 마지막 클릭 이후 250ms 경과 확인"),
    ("3", "checkSession",    "guard:checkSession 호출 — userId:eventId 당 1회 제한 확인"),
    ("✓", "정상 통과",        "antiAbuse.js 승인 → 실제 예매 로직 실행"),
    ("✗", "봇/중복 차단",    "ok: false → 경고 메시지 표시, 예매 차단"),
]
add_table(
    ["결과", "동작", "세부 설명"],
    flow3,
    col_widths=[0.6, 2.5, 8.9],
    header_color=RED
)
doc.add_paragraph()

# ════════════════════════════════════════════════════════════════════
# 5. ERD
# ════════════════════════════════════════════════════════════════════
heading("05  ERD — 핵심 데이터 구조", level=1)

entities = [
    ("AppState",
     "앱 전체 상태",
     "currentUrl: string\ncurrentStep: number\nsteps: Step[]\ntheme: SiteTheme | null\nfontSize: number\nttsEnabled: boolean\nmode: 'slow' | 'orig'"),
    ("Step",
     "단계 정의",
     "id: string\nlabel: string\ntype: 'select' | 'confirm' | 'done'\noptions: string[]\nhint: string\nnextStep: string | null"),
    ("SiteTheme",
     "사이트 테마",
     "primary: string (hex color)\naccent: string (hex color)\nbg: string (hex color)\nname: string\nitalic: boolean"),
    ("GuardStore",
     "어뷰징 방지",
     "clickTimes: Map<sessionId, number[]>\nsessionAttempts: Map<userId:eventId, number>\nHUMAN_MIN_INTERVAL_MS: 250\nSUSPICIOUS_BURST_COUNT: 3\nMAX_ATTEMPTS_PER_EVENT: 1"),
    ("ParsedResult",
     "파싱 결과",
     "steps: Step[]\ntheme: SiteTheme | null\ntitle: string\nsiteKey: string"),
    ("BrowserBounds",
     "BrowserView 위치",
     "x: number (항상 0)\ny: number (navbar.getBoundingClientRect().bottom)\nwidth: number (window.innerWidth)\nheight: number (innerHeight - y)"),
]

for name, role, fields in entities:
    p = doc.add_paragraph()
    set_para_bg(p, NAVY)
    run = p.add_run(f"  {name}  —  {role}")
    run.font.bold = True; run.font.size = Pt(12); run.font.color.rgb = WHITE
    p.paragraph_format.space_before = Pt(8)
    p.paragraph_format.space_after  = Pt(2)

    p2 = doc.add_paragraph()
    set_para_bg(p2, RGBColor(0xF0,0xF4,0xFF))
    run2 = p2.add_run(fields)
    run2.font.name = 'Courier New'; run2.font.size = Pt(9)
    run2.font.color.rgb = NAVY
    p2.paragraph_format.left_indent = Inches(0.3)
    p2.paragraph_format.space_after = Pt(6)

doc.add_paragraph()

# ════════════════════════════════════════════════════════════════════
# 6. IPC API 명세
# ════════════════════════════════════════════════════════════════════
heading("06  IPC API 명세", level=1)

heading("핵심 채널 (Renderer → Main)", level=2, size=14)
core_apis = [
    ("page:load",          "invoke",  "url: string",            "{ steps: Step[], theme: SiteTheme|null, title: string }"),
    ("guard:checkClick",   "invoke",  "sessionId: string",      "{ ok: boolean, reason?: string }"),
    ("guard:checkSession", "invoke",  "userId: string, eventId: string", "{ ok: boolean, reason?: string }"),
    ("tts:speak",          "send",    "text: string",           "void"),
    ("window:set-icon",    "send",    "dataUrl: string (base64 PNG)", "void"),
    ("browser:open",       "invoke",  "url: string",            "{ ok: boolean }"),
    ("browser:close",      "invoke",  "—",                      "{ ok: boolean }"),
    ("browser:navigate",   "invoke",  "url: string",            "{ ok: boolean }"),
    ("browser:back",       "invoke",  "—",                      "{ ok: boolean }"),
    ("browser:forward",    "invoke",  "—",                      "{ ok: boolean }"),
    ("browser:reload",     "invoke",  "—",                      "{ ok: boolean }"),
    ("browser:setBounds",  "invoke",  "{ x, y, width, height }: BrowserBounds", "{ ok: boolean }"),
]
add_table(
    ["채널", "방식", "파라미터", "반환값"],
    core_apis,
    col_widths=[2.5, 1.0, 4.0, 4.5],
    header_color=NAVY
)

heading("Push 이벤트 (Main → Renderer)", level=2, size=14)
push_apis = [
    ("browser:url-changed",   "event send",  "url: string",           "bv-url-display 실시간 업데이트"),
    ("browser:title-changed", "event send",  "title: string",         "윈도우 타이틀 업데이트"),
    ("browser:loading",       "event send",  "isLoading: boolean",    "bv-loading 스피너 표시/숨김"),
]
add_table(
    ["채널", "방식", "페이로드", "설명"],
    push_apis,
    col_widths=[2.5, 1.5, 3.0, 5.0],
    header_color=GREEN
)

# ════════════════════════════════════════════════════════════════════
# 마무리
# ════════════════════════════════════════════════════════════════════
doc.add_paragraph()
p = doc.add_paragraph()
set_para_bg(p, NAVY)
run = p.add_run("  🐢  SlowBro — 핵심 가치 요약")
run.font.size = Pt(16); run.font.bold = True; run.font.color.rgb = WHITE
p.paragraph_format.space_before = Pt(10)
p.paragraph_format.space_after  = Pt(4)

summary = [
    "어르신 접근성:  복잡한 티켓 사이트를 3~4 단계 가이드로 단순화",
    "브랜드 계승:  원본 사이트 색상(SITE_THEMES)으로 신뢰감 유지",
    "이중 모드:  SlowBro 변환 + 원본 BrowserView 병행 제공",
    "안전 예매:  250ms 클릭 가드 + 세션당 1회 예매 제한 (antiAbuse.js)",
    "음성 안내:  Web Speech API ko-KR TTS로 각 단계 읽어주기",
]
for s in summary:
    bullet(s)

# ── 저장 ─────────────────────────────────────────────────────────
out_path = r"C:\khu\khuton\slowbro\demo\SlowBro_기술문서.docx"
doc.save(out_path)
print(f"Saved: {out_path}")
